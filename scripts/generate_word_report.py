import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Mengatur warna latar belakang cell tabel"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Mengatur margin (padding) cell tabel"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    """Membuat heading dengan gaya warna dan ukuran kustom menggunakan Times New Roman"""
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Navy
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x4B, 0x72, 0x9F)  # Steel Blue
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Dark Grey
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
    return p

def main():
    print("[START] Membuat dokumen laporan tugas Word (.docx) dengan font Times New Roman dan 8 screenshot...")
    
    doc = Document()
    
    # 1. Atur Margin Halaman (Standar 1 inci / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default font style ke Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)  # Standar Times New Roman adalah 12 pt untuk laporan akademis
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Charcoal
    
    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_run = title_p.add_run("LAPORAN TUGAS KELOMPOK 2")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x4B, 0x72, 0x9F)
    
    title_main_p = doc.add_paragraph()
    title_main_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_main_p.paragraph_format.space_before = Pt(10)
    title_main_p.paragraph_format.space_after = Pt(10)
    title_main_run = title_main_p.add_run("SHORTEST PATH & MULTI-OBJECTIVE ROUTING FOR SAFAR MODE\n(Aplikasi Web iMosque SafarRoute)")
    title_main_run.font.name = 'Times New Roman'
    title_main_run.font.size = Pt(24)
    title_main_run.font.bold = True
    title_main_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Implementasi Algoritma Shortest Path (Dijkstra & A*) Lintas Jaringan Jalan OpenStreetMap dengan Optimasi Multi-Objective Salat, Kapasitas, dan Biaya Operasional")
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph().paragraph_format.space_before = Pt(80)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_run1 = info_p.add_run("Disusun Oleh:\n")
    info_run1.font.name = 'Times New Roman'
    info_run1.font.bold = True
    
    info_run2 = info_p.add_run("Kelompok 2\n(Placeholder Nama Anggota Kelompok)\n\n")
    info_run2.font.name = 'Times New Roman'
    
    info_run3 = info_p.add_run("Dosen Pengampu:\n(Placeholder Nama Dosen)\n\n")
    info_run3.font.name = 'Times New Roman'
    
    info_run4 = info_p.add_run("PROGRAM STUDI TEKNIK INFORMATIKA\nTahun Akademik 2026")
    info_run4.font.name = 'Times New Roman'
    info_run4.font.bold = True
    
    doc.add_page_break()
    
    # ==================== BAB 1 ====================
    add_heading_styled(doc, "BAB 1: PENDAHULUAN", level=1)
    
    p = doc.add_paragraph(
        "Kebutuhan akan navigasi yang cerdas dan efisien menjadi sangat penting dalam mendukung mobilitas modern. "
        "Bagi umat Muslim yang sedang melakukan perjalanan jauh (safar), pencarian rute tidak hanya terbatas pada "
        "menemukan jalur tercepat dari titik asal ke tujuan. Faktor spiritual seperti waktu pelaksanaan salat fardhu, "
        "ketersediaan fasilitas ibadah (masjid), kapasitas tampung masjid, serta efisiensi biaya perjalanan (seperti konsumsi "
        "bahan bakar minyak dan biaya tol) merupakan variabel krusial yang mempengaruhi kenyamanan perjalanan."
    )
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    
    p = doc.add_paragraph(
        "Aplikasi iMosque SafarRoute dirancang untuk mengatasi permasalahan tersebut dengan menyediakan sistem rekomendasi rute "
        "multi-objective yang cost-efficient tanpa ketergantungan pada Google Maps API yang berbayar. Aplikasi mengintegrasikan data "
        "jaringan jalan OpenStreetMap (OSM) dan data masjid lokal untuk memberikan rekomendasi rute perjalanan safar yang optimal."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "1.1 Rumusan Masalah", level=2)
    p = doc.add_paragraph(
        "1. Bagaimana membangun API routing yang mandiri dan cost-efficient tanpa menggunakan layanan berbayar pihak ketiga (seperti Google Maps)?\n"
        "2. Bagaimana menerapkan algoritma pencarian rute terpendek (shortest path) seperti Dijkstra dan A* pada grafik jaringan jalan yang besar?\n"
        "3. Bagaimana memodelkan optimasi multi-objective yang menggabungkan waktu tempuh, jadwal adzan offline, kapasitas masjid, rating, dan estimasi biaya operasional perjalanan?\n"
        "4. Bagaimana kinerja perbandingan algoritma Dijkstra dan A* dari segi kecepatan komputasi dan penggunaan memori pada data berskala besar?"
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "1.2 Tujuan Proyek", level=2)
    p = doc.add_paragraph(
        "1. Mengimplementasikan algoritma Dijkstra Bidirectional dan A* pada grafik jalan OpenStreetMap yang diunduh secara lokal.\n"
        "2. Membangun REST API berbasis FastAPI yang mampu melakukan kalkulasi rute multi-objective safar secara dinamis.\n"
        "3. Membuat dashboard web interaktif berbasis Next.js dan Leaflet untuk memetakan rute dan lokasi masjid terdekat.\n"
        "4. Mengevaluasi secara mendalam performa algoritma pencarian rute pada subset data besar (wilayah DKI Jakarta) guna memberikan rekomendasi yang akurat dan dapat dipertanggungjawabkan."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== BAB 2 ====================
    add_heading_styled(doc, "BAB 2: ARSITEKTUR SISTEM & ALUR DATA", level=1)
    p = doc.add_paragraph(
        "Sistem iMosque SafarRoute menggunakan arsitektur modern berkinerja tinggi yang memisahkan antara frontend interaktif, "
        "backend komputasional, dan penyimpanan database NoSQL multi-model. Berikut adalah rincian komponen arsitektur sistem:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "2.1 Komponen Sistem", level=2)
    p = doc.add_paragraph(
        "- Frontend (Next.js 16 / React 19): Menyediakan antarmuka pengguna berbasis peta Leaflet yang responsif, visualisasi grafik "
        "perbandingan performa menggunakan Recharts, serta sinkronisasi pengaturan pengguna menggunakan Zustand.\n"
        "- Backend (FastAPI / Python): Mengelola logika bisnis, pembersihan dataset (pandas/scikit-learn), perhitungan waktu salat "
        "offline nasional, Snapping koordinat GPS ke jalan melalui spatial index STRtree, serta eksekusi algoritma routing.\n"
        "- Database (ArangoDB): Menyimpan data masjid, dokumen dataset, metadata grafik cache, dan preferensi pengguna. ArangoDB dipilih "
        "karena keunggulannya dalam mendukung query geospasial (GEO_DISTANCE) secara native pada indeks 2D untuk pencarian masjid terdekat.\n"
        "- Jaringan Jalan (OSM/GraphML): Menggunakan pustaka OSMnx untuk mengunduh grafik jalan dari OpenStreetMap, lalu menyimpannya dalam "
        "format GraphML persisten di data/osm_cache/."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "2.2 Alur Data Utama", level=2)
    p = doc.add_paragraph(
        "Alur data dalam aplikasi iMosque SafarRoute berjalan sebagai berikut:\n"
        "1. Pengunggahan Dataset: Berkas CSV masjid diunggah oleh admin -> diproses melalui pipeline machine learning untuk pembersihan koordinat dan enrichment atribut -> disimpan ke ArangoDB.\n"
        "2. Pemrosesan Grafik OSM: Mengambil bounding box (bbox) wilayah masjid -> mengunduh grafik jalan dari Overpass API -> menyimpan sebagai file GraphML -> mengompilasi menjadi cache biner (*.runtime.pkl dan *.edges.pkl) untuk memangkas cold load time.\n"
        "3. Permintaan Rute Safar: Koordinat GPS awal dan tujuan dikirim oleh frontend -> API mencari kandidat masjid terdekat dalam radius koridor -> Snapping titik awal/akhir ke edge jalan -> Menjalankan Dijkstra/A* -> Mengembalikan rute dalam format compact Google Encoded Polyline 5 dengan kompresi GZip."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== BAB 3 ====================
    add_heading_styled(doc, "BAB 3: IMPLEMENTASI ALGORITMA & MULTI-OBJECTIVE ROUTING", level=1)
    p = doc.add_paragraph(
        "Inti dari sistem SafarRoute adalah kemampuan pencarian rute terpendek yang diintegrasikan dengan logika evaluasi multi-objective "
        "untuk menyaring dan mengurutkan masjid terbaik sepanjang perjalanan."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "3.1 Algoritma Shortest Path Lokal (Tanpa Google Maps API)", level=2)
    p = doc.add_paragraph(
        "Untuk mewujudkan API pencarian rute yang cost-efficient dan berdaulat penuh, aplikasi iMosque tidak menggunakan API Google Maps. "
        "Sebaliknya, sistem membangun grafik jalan OpenStreetMap secara lokal menggunakan OSMnx. Sistem mengimplementasikan dua algoritma pencarian rute terpendek:\n"
        "1. Dijkstra Bidirectional: Melakukan pencarian rute secara simultan dari arah awal (origin) dan akhir (destination) hingga kedua pencarian bertemu di tengah. Algoritma ini secara signifikan mengurangi jumlah node yang perlu dieksplorasi.\n"
        "2. A* (A-Star): Menggunakan fungsi heuristik jarak lurus (great-circle distance) yang konsisten dan admissible untuk mengarahkan pencarian langsung menuju target."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "3.2 Model Skor Multi-Objective", level=2)
    p = doc.add_paragraph(
        "Rekomendasi rute safar diurutkan menggunakan formula multi-objective yang mempertimbangkan 5 komponen utama berikut:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "1. Waktu Tempuh (Travel Time): Durasi berkendara dari asal ke masjid, dan dari masjid ke tujuan akhir.\n"
        "2. Waktu Salat (Adzan Penalty): Penalti waktu jika kedatangan musafir di masjid tidak sinkron dengan waktu salat fardhu terdekat (kalkulasi offline berbasis koordinat geospasial).\n"
        "3. Kapasitas Masjid (Capacity Penalty): Penalti jika kapasitas daya tampung masjid terlalu kecil atau penuh untuk menghindari kerumunan.\n"
        "4. Rating Masjid (Priority Score): Skor prioritas masjid yang dihitung dari bobot ulasan, ketersediaan fasilitas (tempat wudhu, toilet, area parkir, akses difabel), dan rating masjid.\n"
        "5. Estimasi Biaya Rupiah (Financial Cost): Menghitung biaya perjalanan riil berdasarkan konsumsi bahan bakar, biaya operasional kendaraan, dan tarif tol."
    )
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Perhitungan estimasi biaya rupiah dimodelkan menggunakan rumus:\n"
        "Biaya = (Jarak / Efisiensi BBM x Harga BBM) + (Jarak x Biaya Operasional) + (Jarak Tol x Tarif Tol/km)\n\n"
        "Parameter default yang digunakan:\n"
        "- Harga BBM: Rp10.000,- / liter\n"
        "- Efisiensi BBM: 12 km / liter\n"
        "- Biaya Operasional: Rp300,- / km\n"
        "- Tarif Tol rata-rata: Rp1.000,- / km"
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== BAB 4 ====================
    add_heading_styled(doc, "BAB 4: EVALUASI PERFORMA & BENCHMARK ALGORITMA", level=1)
    p = doc.add_paragraph(
        "Evaluasi dilakukan secara ilmiah untuk mengukur performa nyata dari algoritma Dijkstra Bidirectional dan A* pada "
        "subset data besar (Graph DKI Jakarta) yang terdiri atas 196.565 node dan 456.186 edge. Pengujian dijalankan pada "
        "20 pasangan origin-destination (OD) acak dengan jarak lurus antara 3 km hingga 35 km menggunakan seed tetap 20260718."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "4.1 Tabel Hasil Perbandingan Algoritma", level=2)
    
    # Buat Tabel
    table = doc.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Metrik", "Dijkstra Bidirectional", "A* (A-Star)"]
    
    # Format Header
    hdr_cells = table.rows[0].cells
    for i, title_cell in enumerate(headers):
        hdr_cells[i].text = title_cell
        set_cell_background(hdr_cells[i], "1B365D")  # Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p_cell = hdr_cells[i].paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cell.runs[0]
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    data = [
        ["Waktu Rata-rata", "348,88 ms", "581,24 ms"],
        ["Waktu Median", "328,87 ms", "548,95 ms"],
        ["Waktu p95", "737,69 ms", "1.192,82 ms"],
        ["Memori Median", "30.049,41 KB", "11.097,89 KB"],
        ["Expanded Nodes Rata-rata", "35.806,75", "43.156,25"],
        ["Examined Edges Rata-rata", "83.175,25", "99.296,85"]
    ]
    
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p_cell = row_cells[col_idx].paragraphs[0]
            if len(p_cell.runs) > 0:
                p_cell.runs[0].font.name = 'Times New Roman'
            if col_idx > 0:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    p = doc.add_paragraph(
        "Keterangan Hasil:\n"
        "1. Biaya Shortest Path (Optimal Cost) 100% cocok (identik) pada seluruh 20 pasangan OD yang diuji.\n"
        "2. Dijkstra Bidirectional terbukti 30-40% lebih cepat dibandingkan A* satu arah dalam workload pengujian ini.\n"
        "3. A* memiliki kelebihan yang signifikan dalam konsumsi memori median, yaitu sekitar 63% lebih rendah dibandingkan Dijkstra."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "4.2 Validasi Akurasi terhadap OSRM", level=2)
    p = doc.add_paragraph(
        "Sebagai bentuk evaluasi independen, rute lokal yang dihasilkan divalidasi terhadap Open Source Routing Machine (OSRM) publik. "
        "Hasil perbandingan 10 pasangan rute menunjukkan tingkat penyimpangan yang sangat wajar:\n"
        "- Jarak Rute (MAPE): 6,35%\n"
        "- Durasi Rute (MAPE): 12,98%\n\n"
        "Penyimpangan kecil ini wajar terjadi karena adanya perbedaan snaps koordinat GPS ke jalan, pembaruan data jalan real-time pada OSRM, "
        "serta profil kecepatan kendaraan yang diasumsikan oleh masing-masing mesin routing."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== BAB 5 ====================
    add_heading_styled(doc, "BAB 5: DOKUMENTASI ANTARMUKA APLIKASI", level=1)
    p = doc.add_paragraph(
        "Berikut merupakan dokumentasi tangkapan layar (screenshot) antarmuka iMosque SafarRoute dalam tema terang (light mode) "
        "yang menunjukkan pemenuhan seluruh tugas utama (Dijkstra, A*, Graph Masjid, Jaringan Jalan) dan optimasi multi-objective (waktu, adzan, kapasitas, biaya):"
    )
    p.paragraph_format.space_after = Pt(8)
    
    screenshot_dir = os.path.join("outputs", "screenshots")
    
    screenshots = [
        {
            "file": "dashboard_peta.png",
            "title": "Gambar 5.1: Dashboard Peta Utama & Penanda Lokasi Masjid DKI Jakarta",
            "desc": "Tangkapan layar ini menampilkan halaman utama dashboard Next.js dengan visualisasi peta Leaflet. Peta secara akurat menggambarkan penanda lokasi masjid-masjid di Jakarta serta sidebar kosong di sebelah kiri sebelum pengguna menentukan tujuan rute."
        },
        {
            "file": "dashboard_settings.png",
            "title": "Gambar 5.2: Antarmuka Konfigurasi Parameter Optimasi Multi-Objective",
            "desc": "Tangkapan layar ini menampilkan panel popover 'Konfigurasi Rute' (Pengaturan Pencarian) pada Dashboard Utama. Panel ini memvalidasi implementasi parameter multi-objective secara langsung di UI, yang mencakup: pemilihan algoritma (Dijkstra vs A* Heuristik), pilihan profil rute (Balanced, Fastest, Low Cost, Prayer Priority), konfigurasi parameter biaya riil (harga BBM, efisiensi konsumsi bahan bakar, biaya operasional kendaraan, dan tarif tol), penentuan waktu keberangkatan, serta target waktu salat fardhu (adzan)."
        },
        {
            "file": "dashboard_detail_masjid.png",
            "title": "Gambar 5.3: Laci Detail Informasi Masjid Persinggahan Safar",
            "desc": "Tangkapan layar ini menampilkan laci (drawer) detail informasi ketika pengguna mengeklik salah satu marker masjid di peta, dalam contoh ini 'Musholla Wanita Stasiun Gambir'. Drawer menyajikan rating bintang, jumlah ulasan, jarak real-time dari posisi musafir, check-in kunjungan, serta daftar fasilitas masjid terprediksi secara cerdas (seperti AC dan Perpustakaan), lengkap dengan tombol 'Mulai Rute'."
        },
        {
            "file": "dashboard_rute_aktif.png",
            "title": "Gambar 5.4: Jalur Rute Navigasi Teroptimal Tergambar & Rincian Estimasi Biaya Perjalanan",
            "desc": "Tangkapan layar ini menampilkan visualisasi rute navigasi teroptimal terhitung (polyline biru) di peta Leaflet dari Monas ke Stasiun Gambir setelah pengguna mengeklik 'Mulai Rute'. Panel RouteResultPanel di sebelah kiri menyajikan rincian estimasi biaya riil rupiah (pecahan biaya BBM dan operasional), status keamanan waktu kedatangan salat (aman sebelum waktu adzan Maghrib tiba), kapasitas tampung masjid, serta tombol evaluasi cepat untuk membandingkan Dijkstra vs A* secara instan."
        },
        {
            "file": "admin_dashboard.png",
            "title": "Gambar 5.5: Overview Dashboard Admin & Status Integrasi ArangoDB",
            "desc": "Halaman ini menunjukkan tab Dashboard Overview di admin panel. Tab ini menyajikan ringkasan dataset aktif, status server, dan total jumlah dokumen masjid yang terindeks secara geospasial pada database ArangoDB."
        },
        {
            "file": "admin_dataset.png",
            "title": "Gambar 5.6: Pengelolaan Dataset Masjid & Machine Learning Pipeline",
            "desc": "Menampilkan antarmuka pengelolaan berkas CSV masjid (upload, proses ML pipeline, dan setel dataset aktif). Halaman ini bertugas melakukan enrichment data secara dinamis tanpa restart backend."
        },
        {
            "file": "admin_benchmark.png",
            "title": "Gambar 5.7: Panel Uji Performa Komparatif Real-Time (Dijkstra vs A*) dengan Grafik",
            "desc": "Tangkapan layar ini menampilkan hasil eksekusi pengujian performa komparatif antara Dijkstra Bidirectional dan A* secara langsung di UI. Setelah tombol 'Mulai Bandingkan Algoritma' diklik, sistem mengeksekusi routing pada graph jalan DKI Jakarta yang besar secara lokal dan merender visualisasi tabel performa serta grafik batang (bar chart) perbandingan durasi komputasi milidetik, expanded nodes, dan examined edges secara visual."
        },
        {
            "file": "swagger_api_docs.png",
            "title": "Gambar 5.8: Dokumentasi Endpoint REST API FastAPI Backend (Swagger UI)",
            "desc": "Dokumentasi interaktif Swagger UI untuk REST API backend iMosque yang berjalan di port 8000. Dokumentasi ini merangkum seluruh endpoint routing cost-efficient mandiri (tanpa Google Maps API) seperti endpoint /routes/recommend dan /routes/benchmark yang siap diintegrasikan secara luas."
        }
    ]
    
    for img_info in screenshots:
        img_path = os.path.join(screenshot_dir, img_info["file"])
        if os.path.exists(img_path):
            try:
                # 1. Judul Gambar
                p_title = doc.add_paragraph()
                run_title = p_title.add_run(img_info["title"])
                run_title.font.name = 'Times New Roman'
                run_title.font.size = Pt(11)
                run_title.font.bold = True
                p_title.paragraph_format.space_before = Pt(12)
                p_title.paragraph_format.space_after = Pt(4)
                
                # 2. Gambar
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(5.5))
                
                # 3. Keterangan / Deskripsi Gambar (Caption)
                p_cap = doc.add_paragraph()
                run_cap = p_cap.add_run(f"Keterangan: {img_info['desc']}")
                run_cap.font.name = 'Times New Roman'
                run_cap.font.size = Pt(10)
                run_cap.italic = True
                p_cap.paragraph_format.space_after = Pt(14)
                
            except Exception as e:
                p_err = doc.add_paragraph(f"[Gagal memuat gambar {img_info['file']}: {e}]")
                p_err.runs[0].font.name = 'Times New Roman'
                p_err.italic = True
        else:
            p_missing = doc.add_paragraph(f"[Gambar tidak ditemukan di: {img_path}]")
            p_missing.runs[0].font.name = 'Times New Roman'
            p_missing.italic = True
            
    doc.add_page_break()
    
    # ==================== BAB 6 ====================
    add_heading_styled(doc, "BAB 6: KESIMPULAN & REKOMENDASI", level=1)
    
    add_heading_styled(doc, "6.1 Kesimpulan", level=2)
    p = doc.add_paragraph(
        "1. Aplikasi iMosque SafarRoute berhasil menyajikan rute multi-objective safar yang akurat dan berkinerja tinggi "
        "dengan memanfaatkan grafik jalan lokal OpenStreetMap dan database geospasial ArangoDB tanpa dependensi Google Maps.\n"
        "2. Algoritma Dijkstra Bidirectional merupakan pilihan default tercepat untuk mesin routing lokal ini (median 328,87 ms), "
        "sementara algoritma A* memberikan alternatif hemat memori yang sangat baik (hemat 63% memori).\n"
        "3. Integrasi model multi-objective (waktu salat offline, kapasitas masjid, priority score, dan biaya rupiah perjalanan) "
        "mampu memberikan rekomendasi masjid persinggahan safar yang presisi dan relevan dengan preferensi pengguna."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_heading_styled(doc, "6.2 Rekomendasi Pengembangan", level=2)
    p = doc.add_paragraph(
        "1. Skalabilitas Multi-Region: Perluasan area grafik jalan ke tingkat nasional menggunakan server routing terdistribusi "
        "(routing worker regional) agar tidak membebani penggunaan RAM server backend utama.\n"
        "2. Integrasi Data Real-Time: Integrasi data kemacetan lalu lintas, laporan cuaca secara real-time, serta informasi pemeliharaan "
        "masjid secara berkala untuk meningkatkan validitas estimasi waktu tempuh.\n"
        "3. Sistem Keamanan & Autentikasi: Penambahan layer otorisasi dan otentikasi admin (seperti JWT token) pada endpoint administrasi "
        "upload dataset, manipulasi data masjid, dan build graph sebelum sistem dideploy di lingkungan produksi."
    )
    p.paragraph_format.space_after = Pt(24)
    
    # Simpan dokumen
    output_filename = "Laporan_Tugas_Kelompok_2_SafarRoute.docx"
    doc.save(output_filename)
    print(f"[SUCCESS] File laporan berhasil disimpan sebagai: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    main()
